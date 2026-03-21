#!/usr/bin/env python
"""Create CodeNest project documentation in DOCX format"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_centered_text(text, size=12, bold=False):
    """Add centered text"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(size)
    run.font.bold = bold
    return para

def add_heading_custom(text, level=1):
    """Add custom heading"""
    heading = doc.add_heading(text, level=level)
    return heading

# Title Page
add_centered_text('A', 16, True)
add_centered_text('Project Report', 16, True)
add_centered_text('On', 14)
add_centered_text('"CODENEST - AI-POWERED COMPETITIVE PROGRAMMING PLATFORM"', 14, True)
doc.add_paragraph()
add_centered_text('Submitted in partial fulfillment of the requirements for the award of the degree', 12)
add_centered_text('of', 12)
add_centered_text('BACHELOR OF TECHNOLOGY', 12, True)
add_centered_text('IN', 12)
add_centered_text('COMPUTER SCIENCE AND ENGINEERING', 12, True)
doc.add_paragraph()
doc.add_paragraph()
add_centered_text('Submitted By', 12, True)
add_centered_text('[Student Name 1] ([Roll Number])', 12)
add_centered_text('[Student Name 2] ([Roll Number])', 12)
add_centered_text('[Student Name 3] ([Roll Number])', 12)
doc.add_paragraph()
add_centered_text('Under the esteemed guidance of', 12, True)
add_centered_text('[Guide Name], M.Tech., Ph.D.', 12)
add_centered_text('Assistant Professor', 12)
doc.add_paragraph()
doc.add_paragraph()
add_centered_text('Department of Computer Science and Engineering', 12)
add_centered_text('[INSTITUTION NAME]', 12, True)
add_centered_text('(AUTONOMOUS)', 12)
add_centered_text('[Address]', 12)
add_centered_text('2025-2026', 12)

# Page break
doc.add_page_break()

# Certificate Page
add_centered_text('[INSTITUTION NAME]', 14, True)
add_centered_text('(AUTONOMOUS)', 12)
add_centered_text('[Affiliation Details]', 11)
add_centered_text('[Address]', 11)
doc.add_paragraph()
add_centered_text('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', 12, True)
doc.add_paragraph()
doc.add_paragraph()
add_centered_text('Certificate', 14, True)
doc.add_paragraph()

cert_text = 'This is to certify that the Project report entitled "CodeNest - AI-Powered Competitive Programming Platform" is the Bonafide work carried out by [Student Names] bearing Roll Numbers [Roll Numbers] in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science & Engineering during the academic year "2025-2026"'
para = doc.add_paragraph(cert_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Signatures
sig_para = doc.add_paragraph()
sig_para.add_run('Project Guide').bold = True
sig_para.add_run('\t\t\t')
sig_para.add_run('Head of the Department').bold = True

doc.add_paragraph('[Guide Name]')
doc.add_paragraph('Assistant Professor')
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('Date:')
doc.add_paragraph('Place:')
doc.add_paragraph()
add_centered_text('External Examiner', 12)

# Page break
doc.add_page_break()

# Declaration
add_centered_text('DECLARATION CERTIFICATE', 14, True)
doc.add_paragraph()

decl_text = 'We students of Computer Science & Engineering, [INSTITUTION NAME] (AUTONOMOUS), hereby declare that the dissertation entitled "CodeNest - AI-Powered Competitive Programming Platform" embodies the report of our project work carried out by us during IV year under the guidance of [Guide Name], Assistant Professor, Department of Computer Science & Engineering, and this work has been submitted for the partial fulfillment of the requirements for the award of degree of Bachelor of Technology. The results embodied in this project report have not been submitted to any other University or Institute for the award of any Degree or Diploma.'
para = doc.add_paragraph(decl_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph('Date:')
doc.add_paragraph('Place:')
doc.add_paragraph()
doc.add_paragraph('[Student Signatures]')

# Page break
doc.add_page_break()

# Acknowledgement
add_centered_text('ACKNOWLEDGEMENT', 14, True)
doc.add_paragraph()

ack_text = '''The successful completion of this project would be incomplete without acknowledging those whose constant support and guidance made it possible. We are grateful for the opportunity to express our thanks to all of them.

We express our sincere gratitude to our guide & Head of the Department, [Guide Name], Assistant Professor, Computer Science & Engineering, for continuous guidance, encouragement, and constructive feedback at every stage of the project, which greatly contributed to its successful completion.

We express our deepfelt gratitude to the Project Coordinators for their valuable guidance and unstinting encouragement that enabled us to accomplish our project successfully in time.

We extend our special thanks to [Principal Name], Principal of [Institution Name], for providing the necessary information for our project. We also thank the faculty, non-teaching staff, and friends who directly or indirectly supported us in completing it on time.

We sincerely thank the Management for the excellent facilities and our families for their constant support throughout the project.'''

para = doc.add_paragraph(ack_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
doc.add_paragraph('Project Associates')
doc.add_paragraph('[Student Name 1] ([Roll Number])')
doc.add_paragraph('[Student Name 2] ([Roll Number])')
doc.add_paragraph('[Student Name 3] ([Roll Number])')

# Page break
doc.add_page_break()

# Abstract
add_centered_text('ABSTRACT', 14, True)
doc.add_paragraph()

abstract_text = '''Competitive programming has emerged as a critical skill development activity for computer science students, fostering algorithmic thinking, problem-solving abilities, and coding proficiency. However, existing platforms often lack comprehensive features for educational institutions, including mentor-student interactions, real-time code execution, AI-powered assistance, and integrated contest management. This project proposes CodeNest, an AI-Powered Competitive Programming Platform designed specifically for educational environments to bridge these gaps.

CodeNest is a full-stack web application built using React.js for the frontend and Django REST Framework for the backend, with PostgreSQL/SQLite for data management. The platform supports three primary user roles: Students, Teachers (Mentors), and Administrators, each with role-specific functionalities and dashboards. Students can solve problems across multiple difficulty levels and topics, participate in timed contests, track their progress through comprehensive analytics, and receive AI-powered hints and explanations. Teachers can create and manage contests, monitor student performance, assign problems, and access class-wide analytics to identify learning gaps.

The system implements a secure Docker-based code execution environment supporting multiple programming languages including Python, Java, C++, and JavaScript. An integrated AI chatbot powered by the Gemini API provides contextual assistance, explains problem concepts, and offers debugging support without revealing complete solutions. The platform features real-time leaderboards, achievement systems, discussion forums for peer learning, and activity heatmaps visualizing student engagement patterns.

CodeNest addresses critical limitations in existing platforms by providing institutional-grade features including batch management, branch-wise analytics, mentor-student assignment, and customizable problem sets. The platform's modular architecture ensures scalability, maintainability, and extensibility for future enhancements. Comprehensive testing demonstrates the system's effectiveness in improving student engagement, learning outcomes, and competitive programming skills.

Keywords: Competitive Programming, AI-Powered Learning, Code Execution, Contest Management, Educational Platform, Django, React, Docker, Real-time Analytics, Mentor Dashboard.'''

para = doc.add_paragraph(abstract_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Page break
doc.add_page_break()

# Table of Contents
add_centered_text('TABLE OF CONTENTS', 14, True)
doc.add_paragraph()

toc_items = [
    ('Certificate', 'ii'),
    ('Declaration', 'iii'),
    ('Acknowledgement', 'iv'),
    ('Abstract', 'v'),
    ('List of Figures', 'viii'),
    ('List of Tables', 'ix'),
    ('List of Abbreviations', 'x'),
    ('', ''),
    ('CHAPTER 1: INTRODUCTION', '1'),
    ('1.1 Background and Motivation', '1'),
    ('1.2 Problem Statement', '3'),
    ('1.3 Objectives of the Study', '5'),
    ('1.4 Scope of the Project', '6'),
    ('1.5 Organization of the Thesis', '8'),
    ('', ''),
    ('CHAPTER 2: LITERATURE SURVEY', '9'),
    ('2.1 Overview of Competitive Programming Platforms', '9'),
    ('2.2 AI in Educational Technology', '12'),
    ('2.3 Code Execution and Security', '15'),
    ('2.4 Contest Management Systems', '18'),
    ('2.5 Analytics and Progress Tracking', '20'),
    ('', ''),
    ('CHAPTER 3: SYSTEM ANALYSIS', '23'),
    ('3.1 Existing System Limitations', '23'),
    ('3.2 Proposed System Features', '25'),
    ('3.3 Benefits and Advantages', '28'),
    ('', ''),
    ('CHAPTER 4: REQUIREMENTS ANALYSIS', '30'),
    ('4.1 Functional Requirements', '30'),
    ('4.2 Non-Functional Requirements', '33'),
    ('4.3 Hardware Requirements', '35'),
    ('4.4 Software Requirements', '36'),
    ('', ''),
    ('CHAPTER 5: SYSTEM DESIGN', '38'),
    ('5.1 System Architecture', '38'),
    ('5.2 Database Design', '41'),
    ('5.3 Module Design', '45'),
    ('5.4 User Interface Design', '50'),
    ('', ''),
    ('CHAPTER 6: IMPLEMENTATION', '55'),
    ('6.1 Technology Stack', '55'),
    ('6.2 Backend Implementation', '58'),
    ('6.3 Frontend Implementation', '62'),
    ('6.4 AI Integration', '66'),
    ('6.5 Code Execution Engine', '69'),
    ('6.6 Security Implementation', '72'),
    ('', ''),
    ('CHAPTER 7: RESULTS AND DISCUSSION', '75'),
    ('7.1 System Performance Analysis', '75'),
    ('7.2 User Testing and Feedback', '78'),
    ('7.3 Comparative Analysis', '81'),
    ('7.4 Achievement System Effectiveness', '84'),
    ('', ''),
    ('CHAPTER 8: CHALLENGES AND SOLUTIONS', '87'),
    ('8.1 Technical Challenges', '87'),
    ('8.2 Implementation Challenges', '89'),
    ('8.3 Solutions and Workarounds', '91'),
    ('', ''),
    ('CHAPTER 9: CONCLUSION AND FUTURE SCOPE', '94'),
    ('9.1 Summary of Work', '94'),
    ('9.2 Contributions', '95'),
    ('9.3 Future Enhancements', '96'),
    ('9.4 Conclusion', '98'),
    ('', ''),
    ('REFERENCES', '100'),
    ('APPENDICES', '105'),
]

for item, page in toc_items:
    if item == '':
        doc.add_paragraph()
    else:
        para = doc.add_paragraph()
        para.add_run(item)
        para.add_run('\t' * 5)
        para.add_run(page)

# Page break
doc.add_page_break()

# Chapter 1: Introduction
doc.add_heading('CHAPTER 1', 1)
doc.add_heading('INTRODUCTION', 1)
doc.add_paragraph()

intro_text = '''Competitive programming has become an essential component of computer science education, serving as a bridge between theoretical knowledge and practical problem-solving skills. It challenges students to think algorithmically, optimize solutions, and write efficient code under time constraints. Major technology companies including Google, Facebook, Amazon, and Microsoft actively recruit candidates based on their competitive programming performance, making it a critical skill for career advancement in the software industry.

Educational institutions across India have recognized the importance of competitive programming in developing well-rounded computer science professionals. However, the integration of competitive programming into institutional curricula faces several challenges. Existing platforms such as LeetCode, HackerRank, and Codeforces, while excellent for individual practice, lack features specifically designed for educational environments. These platforms do not provide institutional management capabilities, mentor-student interactions, batch-wise analytics, or customizable problem sets aligned with academic curricula.

The growing demand for intelligent, scalable, and institution-friendly competitive programming platforms has driven the development of CodeNest. This project proposes a comprehensive AI-powered platform that combines the best features of existing competitive programming websites with educational management capabilities, AI-assisted learning, and institutional analytics. The system serves three primary user groups - Students, Teachers (Mentors), and Administrators - each with distinct roles and functionalities designed to support effective learning and teaching.

This chapter provides an overview of the background and motivation behind the project, defines the problem statement, outlines the objectives and scope, and presents the organization of the thesis.'''

para = doc.add_paragraph(intro_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
doc.add_heading('1.1 Background and Motivation', 2)

bg_text = '''The landscape of computer science education has evolved significantly over the past two decades. Traditional teaching methods that focused primarily on theoretical concepts and basic programming syntax are no longer sufficient to prepare students for the demands of modern software development. Employers increasingly seek candidates who can demonstrate strong problem-solving abilities, algorithmic thinking, and the capacity to write efficient, optimized code.

Competitive programming emerged as a response to this need, providing a structured framework for developing these critical skills. Platforms like TopCoder (founded in 2001), Codeforces (2010), and LeetCode (2015) have created global communities of programmers who regularly participate in contests, solve algorithmic challenges, and improve their coding proficiency. Research has consistently shown that students who actively engage in competitive programming demonstrate superior performance in technical interviews, coding assessments, and real-world software development tasks.

Despite the proven benefits of competitive programming, its adoption within Indian educational institutions remains inconsistent and often informal. Many colleges lack dedicated platforms for organizing internal contests, tracking student progress, or providing structured guidance. Students typically practice on external platforms without institutional oversight, making it difficult for faculty to monitor progress, identify struggling students, or align practice activities with curriculum objectives.

The COVID-19 pandemic further highlighted the need for robust digital infrastructure in education. Institutions that relied on physical computer labs and in-person coding sessions found their programs disrupted, while those with digital platforms were able to continue programming education seamlessly. This experience has accelerated the recognition that modern educational institutions require comprehensive digital platforms for all aspects of computer science education, including competitive programming.

The motivation for CodeNest stems from direct observation of these challenges within institutional contexts. By developing an AI-powered, institution-friendly competitive programming platform, this project aims to provide colleges with a modern, scalable solution that combines the best features of existing platforms with educational management capabilities, AI-assisted learning, and comprehensive analytics.'''

para = doc.add_paragraph(bg_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
doc.add_heading('1.2 Problem Statement', 2)

problem_text = '''Traditional approaches to teaching competitive programming in educational institutions face several critical limitations that reduce their effectiveness and scalability:

1. Lack of Institutional Integration: Existing competitive programming platforms are designed for individual users and do not provide features for institutional management. Faculty cannot create private contests for their students, assign specific problems aligned with curriculum topics, or access class-wide analytics to identify learning gaps.

2. Absence of Mentor-Student Interaction: Students practicing on external platforms receive no guidance from their instructors. When students encounter difficulties, they have no direct channel to seek help from faculty, leading to frustration and reduced engagement.

3. No AI-Powered Learning Support: Traditional platforms provide problem statements and test cases but offer no intelligent assistance. Students who get stuck on problems have limited options beyond searching for complete solutions online, which undermines the learning process.

4. Limited Progress Tracking: Faculty have no visibility into student practice activities, making it impossible to monitor progress, identify struggling students, or provide timely interventions. Students themselves lack comprehensive analytics showing their strengths, weaknesses, and improvement trends.

5. Inadequate Contest Management: Organizing internal coding contests requires significant manual effort, including problem selection, test case creation, participant management, and result compilation. Many institutions avoid conducting regular contests due to this administrative burden.

6. Security and Plagiarism Concerns: Code execution on student machines or unsecured servers poses security risks. Additionally, detecting plagiarism in coding contests remains challenging without specialized tools.

7. Lack of Gamification and Engagement: Without achievement systems, leaderboards, and progress visualization, student motivation tends to decline over time, particularly for those who do not see immediate improvement.

To address these challenges, this project proposes CodeNest, a comprehensive AI-powered competitive programming platform specifically designed for educational institutions, providing integrated contest management, AI-assisted learning, mentor dashboards, and comprehensive analytics.'''

para = doc.add_paragraph(problem_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()
doc.add_heading('1.3 Objectives of the Study', 2)

obj_text = '''The primary objective of this study is to design and develop CodeNest, an AI-powered competitive programming platform tailored for educational institutions. The specific objectives include:

1. To develop a secure, role-based web application supporting Students, Teachers, and Administrators with role-specific functionalities and access controls.

2. To implement a Docker-based code execution engine supporting multiple programming languages (Python, Java, C++, JavaScript) with secure sandboxing and resource limits.

3. To integrate an AI-powered chatbot using the Gemini API that provides contextual hints, explains concepts, and assists with debugging without revealing complete solutions.

4. To design and implement a comprehensive contest management system allowing teachers to create timed contests, manage participants, and automatically generate leaderboards.

5. To create interactive analytics dashboards for students showing topic-wise progress, submission history, activity heatmaps, and performance trends.

6. To develop mentor dashboards providing class-wide analytics, student performance tracking, and tools for identifying struggling students.

7. To implement an achievement and gamification system that rewards consistent practice, problem-solving milestones, and contest participation.

8. To design a discussion forum enabling peer-to-peer learning, problem discussions, and community engagement.

9. To ensure system security through JWT-based authentication, role-based access control, secure code execution, and data protection measures.

10. To evaluate the system's effectiveness through user testing, performance analysis, and comparison with existing platforms.'''

para = doc.add_paragraph(obj_text)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Save document
doc.save('codenest.docx')
print("✅ Document created successfully: codenest.docx")
