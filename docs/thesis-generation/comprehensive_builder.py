#!/usr/bin/env python
"""Comprehensive CodeNest Documentation Builder - Matches thesis length"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocBuilder:
    def __init__(self, filename):
        self.doc = Document()
        self.filename = filename
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
    
    def centered(self, text, size=12, bold=False):
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(size)
        run.font.bold = bold
        return para
    
    def justified(self, text):
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return para
    
    def list_item(self, text):
        para = self.doc.add_paragraph(text, style='List Paragraph')
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return para
    
    def heading(self, text, level=1):
        return self.doc.add_heading(text, level=level)
    
    def page_break(self):
        self.doc.add_page_break()
    
    def save(self):
        self.doc.save(self.filename)
        print(f"✅ Saved: {self.filename}")

# Initialize builder
builder = DocBuilder('codenest_comprehensive.docx')

# TITLE PAGE
builder.centered('A', 16, True)
builder.centered('Project Report', 16, True)
builder.centered('On', 14)
builder.centered('"CODENEST - AI-POWERED COMPETITIVE', 14, True)
builder.centered('PROGRAMMING PLATFORM"', 14, True)
builder.doc.add_paragraph()
builder.centered('Submitted in partial fulfillment of the requirements for the award of the degree', 12)
builder.centered('of', 12)
builder.centered('BACHELOR OF TECHNOLOGY', 12, True)
builder.centered('IN', 12)
builder.centered('COMPUTER SCIENCE AND ENGINEERING', 12, True)

builder.save()
print("Step 1/10 complete")
