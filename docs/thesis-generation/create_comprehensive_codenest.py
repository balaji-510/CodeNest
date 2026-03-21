#!/usr/bin/env python
"""Create comprehensive CodeNest documentation matching thesis length"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_centered(text, size=12, bold=False):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(size)
    run.font.bold = bold
    return para

def add_justified(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_list_item(text):
    para = doc.add_paragraph(text, style='List Paragraph')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

# TITLE PAGE
add_centered('A', 16, True)
add_centered('Project Report', 16, True)
add_centered('On', 14)
add_centered('"CODENEST - AI-POWERED COMPETITIVE', 14, True)
add_centered('PROGRAMMING PLATFORM"', 14, True)
doc.add_paragraph()
add_centered('Submitted in partial fulfillment of the requirements for the award of the degree', 12)
add_centered('of', 12)
add_centered('BACHELOR OF TECHNOLOGY', 12, True)
add_centered('IN', 12)
add_centered('COMPUTER SCIENCE AND ENGINEERING', 12, True)

doc.save('codenest_comprehensive.docx')
print("✅ Part 1: Title page created")
