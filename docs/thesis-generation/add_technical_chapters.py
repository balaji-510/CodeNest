#!/usr/bin/env python
"""Add technical chapters to CodeNest documentation"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('codenest_final.docx')

def add_justified(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

# Continue Chapter 1
doc.add_heading('1.3 Objectives', 2)
obj = '''The primary objective is to design and develop CodeNest, an AI-powered competitive programming platform for educational institutions. Specific objectives:

1. Develop secure, role-based web application for Students, Teachers, and Administrators
2. Implement Docker-based code execution supporting Python, Java, C++, JavaScript
3. Integrate AI chatbot using Gemini API for contextual assistance
4. Design comprehensive contest management system
5. Create interactive analytics dashboards
6. Develop mentor dashboards with class-wide analytics
7. Implement achievement and gamification system
8. Design discussion forum for peer learning
9. Ensure security through JWT authentication and role-based access
10. Evaluate effectiveness through user testing and performance analysis'''
add_justified(obj)

doc.add_heading('1.4 Scope', 2)
scope = '''CodeNest encompasses complete competitive programming lifecycle:

Student Features: Problem solving, multi-language editor, real-time execution, contest participation, AI hints, progress tracking, achievements, discussion forums, activity heatmaps

Teacher Features: Contest creation, problem assignment, class analytics, student monitoring, batch management, mentor dashboard, leaderboard access

Administrator Features: User management, system configuration, platform analytics, content moderation, database management, security control

The platform focuses on institutional competitive programming. Future enhancements may include external platform integration and mobile applications.'''
add_justified(scope)

doc.add_page_break()

# Chapter 2
doc.add_heading('CHAPTER 2', 1)
doc.add_heading('LITERATURE SURVEY', 1)
doc.add_paragraph()

lit = '''Competitive programming platforms have evolved significantly. Early platforms like TopCoder (2001) and Codeforces (2010) established the foundation. LeetCode (2015) popularized interview preparation.

Research shows students engaging in competitive programming demonstrate superior technical interview performance. However, existing platforms lack institutional features.

AI in education has shown promise for personalized learning. Intelligent tutoring systems provide adaptive feedback. Code execution security remains critical, with Docker containerization providing isolation.'''
add_justified(lit)

doc.add_heading('2.1 Existing Platforms', 2)
existing = '''LeetCode: Focuses on interview preparation, lacks institutional features
HackerRank: Provides some educational features but limited customization
Codeforces: Excellent for contests but no institutional management
CodeChef: Strong community but no mentor dashboards

Gap: None provide comprehensive institutional management with AI assistance and mentor-student interactions.'''
add_justified(existing)

doc.add_page_break()

# Chapter 3
doc.add_heading('CHAPTER 3', 1)
doc.add_heading('SYSTEM ANALYSIS', 1)
doc.add_paragraph()

doc.add_heading('3.1 Existing System Limitations', 2)
limitations = '''Current approaches face several limitations:

1. No institutional integration or batch management
2. Limited mentor-student interaction channels
3. Absence of AI-powered learning assistance
4. Inadequate progress tracking and analytics
5. Complex contest management requiring manual effort
6. Security concerns in code execution
7. Lack of gamification reducing engagement'''
add_justified(limitations)

doc.add_heading('3.2 Proposed System', 2)
proposed = '''CodeNest addresses these limitations through:

1. Role-based architecture (Student, Teacher, Admin)
2. Docker-based secure code execution
3. AI chatbot integration for learning support
4. Comprehensive analytics dashboards
5. Automated contest management
6. Achievement and gamification system
7. Discussion forums for peer learning
8. Real-time leaderboards and progress tracking'''
add_justified(proposed)

doc.add_page_break()

# Chapter 4
doc.add_heading('CHAPTER 4', 1)
doc.add_heading('REQUIREMENTS ANALYSIS', 1)
doc.add_paragraph()

doc.add_heading('4.1 Functional Requirements', 2)
func = '''User Management: Registration, authentication, role-based access
Problem Management: CRUD operations, test cases, difficulty levels
Code Execution: Multi-language support, test case validation, time/memory limits
Contest Management: Creation, scheduling, participant management, leaderboards
AI Assistant: Context-aware hints, concept explanations, debugging support
Analytics: Student progress, topic mastery, submission history, activity heatmaps
Discussion Forum: Post creation, replies, voting, moderation
Achievement System: Badges, milestones, progress tracking'''
add_justified(func)

doc.add_heading('4.2 Non-Functional Requirements', 2)
nonfunc = '''Performance: Response time < 2 seconds, support 1000+ concurrent users
Security: JWT authentication, encrypted passwords, secure code execution
Scalability: Horizontal scaling capability, database optimization
Usability: Intuitive UI, responsive design, accessibility compliance
Reliability: 99.9% uptime, automated backups, error handling
Maintainability: Modular architecture, comprehensive documentation'''
add_justified(nonfunc)

doc.add_heading('4.3 Technology Stack', 2)
tech = '''Frontend: React.js, Tailwind CSS, Recharts, Monaco Editor
Backend: Django REST Framework, Python
Database: PostgreSQL/SQLite
Code Execution: Docker, Judge0 API
AI Integration: Google Gemini API
Authentication: JWT (JSON Web Tokens)
Deployment: Nginx, Gunicorn, Docker Compose'''
add_justified(tech)

doc.save('codenest_final.docx')
print("✅ Added Chapters 2, 3, 4")
