#!/usr/bin/env python
"""Build complete CodeNest documentation - All chapters"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('codenest_final.docx')

def add_justified(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

def add_centered(text, size=12, bold=False):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(size)
    run.font.bold = bold
    return para

# Acknowledgement
add_centered('ACKNOWLEDGEMENT', 14, True)
doc.add_paragraph()
ack = '''The successful completion of this project would be incomplete without acknowledging those whose constant support and guidance made it possible.

We express our sincere gratitude to our guide, [Guide Name], for continuous guidance, encouragement, and constructive feedback at every stage of the project.

We extend our special thanks to [Principal Name], Principal of [Institution Name], for providing the necessary facilities. We also thank the faculty, staff, and friends who supported us.

We sincerely thank the Management and our families for their constant support throughout the project.'''
add_justified(ack)
doc.add_paragraph()
doc.add_paragraph('Project Associates')
doc.add_paragraph('[Student Names and Roll Numbers]')

doc.add_page_break()

# Abstract
add_centered('ABSTRACT', 14, True)
doc.add_paragraph()
abstract = '''Competitive programming has emerged as a critical skill for computer science students. However, existing platforms lack comprehensive features for educational institutions. This project proposes CodeNest, an AI-Powered Competitive Programming Platform designed for educational environments.

CodeNest is built using React.js and Django REST Framework, supporting Students, Teachers, and Administrators. Students can solve problems, participate in contests, and receive AI-powered assistance. Teachers can create contests, monitor performance, and access analytics.

The system implements secure Docker-based code execution supporting Python, Java, C++, and JavaScript. An integrated AI chatbot powered by Gemini API provides contextual assistance. The platform features real-time leaderboards, achievements, discussion forums, and activity heatmaps.

CodeNest addresses limitations in existing platforms by providing institutional features including batch management, mentor-student assignment, and customizable problem sets.

Keywords: Competitive Programming, AI-Powered Learning, Code Execution, Contest Management, Django, React, Docker, Analytics.'''
add_justified(abstract)

doc.add_page_break()

# Chapter 1
doc.add_heading('CHAPTER 1', 1)
doc.add_heading('INTRODUCTION', 1)
doc.add_paragraph()

intro = '''Competitive programming has become essential in computer science education, bridging theoretical knowledge and practical problem-solving. Major technology companies actively recruit based on competitive programming performance.

Educational institutions face challenges integrating competitive programming into curricula. Existing platforms lack institutional management, mentor-student interactions, and customizable problem sets.

CodeNest addresses these gaps by combining competitive programming features with educational management, AI-assisted learning, and institutional analytics.'''
add_justified(intro)

doc.add_heading('1.1 Background and Motivation', 2)
bg = '''Traditional teaching methods focusing on theory are insufficient for modern software development. Competitive programming provides structured skill development.

Despite proven benefits, adoption in Indian institutions remains inconsistent. The COVID-19 pandemic highlighted the need for robust digital infrastructure.

CodeNest aims to provide colleges with a modern, scalable solution combining competitive programming with educational management and AI-assisted learning.'''
add_justified(bg)

doc.add_heading('1.2 Problem Statement', 2)
problem = '''Traditional competitive programming education faces critical limitations:

1. Lack of Institutional Integration - No features for institutional management
2. Absence of Mentor-Student Interaction - No guidance channels
3. No AI-Powered Learning Support - Limited assistance when stuck
4. Limited Progress Tracking - No visibility into student activities
5. Inadequate Contest Management - High administrative burden
6. Security and Plagiarism Concerns - Execution and detection challenges
7. Lack of Gamification - Declining motivation over time

CodeNest addresses these challenges with integrated contest management, AI assistance, mentor dashboards, and comprehensive analytics.'''
add_justified(problem)

doc.save('codenest_final.docx')
print("✅ Added Abstract and Chapter 1")
