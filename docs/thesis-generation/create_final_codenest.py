#!/usr/bin/env python
"""Create complete CodeNest project documentation"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_centered(text, size=12, bold=False):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(size)
    run.font.bold = bold
    return para

def add_justified(text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return para

# Title Page
add_centered('A', 16, True)
add_centered('Project Report', 16, True)
add_centered('On', 14)
add_centered('"CODENEST - AI-POWERED COMPETITIVE PROGRAMMING PLATFORM"', 14, True)
doc.add_paragraph()
add_centered('Submitted in partial fulfillment of the requirements for the award of the degree', 12)
add_centered('of', 12)
add_centered('BACHELOR OF TECHNOLOGY', 12, True)
add_centered('IN', 12)
add_centered('COMPUTER SCIENCE AND ENGINEERING', 12, True)
doc.add_paragraph()
doc.add_paragraph()
add_centered('Submitted By', 12, True)
add_centered('[Student Name 1] ([Roll Number])', 12)
add_centered('[Student Name 2] ([Roll Number])', 12)
add_centered('[Student Name 3] ([Roll Number])', 12)
doc.add_paragraph()
add_centered('Under the esteemed guidance of', 12, True)
add_centered('[Guide Name], M.Tech., Ph.D.', 12)
add_centered('Assistant Professor', 12)
doc.add_paragraph()
doc.add_paragraph()
add_centered('Department of Computer Science and Engineering', 12)
add_centered('[INSTITUTION NAME]', 12, True)
add_centered('(AUTONOMOUS)', 12)
add_centered('[Address]', 12)
add_centered('2025-2026', 12)

doc.add_page_break()

# Save
doc.save('codenest_final.docx')
print("✅ Created codenest_final.docx - Part 1")

# Continue with more content
doc = Document('codenest_final.docx')

# Certificate Page
add_centered('[INSTITUTION NAME]', 14, True)
add_centered('(AUTONOMOUS)', 12)
add_centered('[Affiliation Details]', 11)
add_centered('[Address]', 11)
doc.add_paragraph()
add_centered('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', 12, True)
doc.add_paragraph()
doc.add_paragraph()
add_centered('Certificate', 14, True)
doc.add_paragraph()

cert_text = 'This is to certify that the Project report entitled "CodeNest - AI-Powered Competitive Programming Platform" is the Bonafide work carried out by [Student Names] bearing Roll Numbers [Roll Numbers] in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science & Engineering during the academic year "2025-2026"'
add_justified(cert_text)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
sig_para = doc.add_paragraph()
sig_para.add_run('Project Guide').bold = True
sig_para.add_run('\t\t\t')
sig_para.add_run('Head of the Department').bold = True
doc.add_paragraph('[Guide Name]')
doc.add_paragraph('Assistant Professor')
doc.add_paragraph()
doc.add_paragraph('Date:')
doc.add_paragraph('Place:')

doc.add_page_break()

# Declaration
add_centered('DECLARATION CERTIFICATE', 14, True)
doc.add_paragraph()
decl_text = 'We students of Computer Science & Engineering, [INSTITUTION NAME] (AUTONOMOUS), hereby declare that the dissertation entitled "CodeNest - AI-Powered Competitive Programming Platform" embodies the report of our project work carried out by us during IV year under the guidance of [Guide Name], Assistant Professor, Department of Computer Science & Engineering, and this work has been submitted for the partial fulfillment of the requirements for the award of degree of Bachelor of Technology.'
add_justified(decl_text)
doc.add_paragraph()
doc.add_paragraph('Date:')
doc.add_paragraph('Place:')

doc.add_page_break()

doc.save('codenest_final.docx')
print("✅ Added Certificate and Declaration")
