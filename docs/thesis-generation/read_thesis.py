#!/usr/bin/env python
"""Read thesis document structure"""
from docx import Document

doc = Document('../thesis_A-14.docx')

print('=== DOCUMENT STRUCTURE ===\n')
for i, para in enumerate(doc.paragraphs[:80]):
    if para.text.strip():
        print(f'{i}: [{para.style.name}] {para.text[:150]}')

print('\n=== HEADINGS ===\n')
for para in doc.paragraphs:
    if 'Heading' in para.style.name and para.text.strip():
        print(f'[{para.style.name}] {para.text}')
