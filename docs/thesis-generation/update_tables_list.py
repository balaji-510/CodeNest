#!/usr/bin/env python
"""Update List of Tables with more relevant entries for CodeNest"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Reviewing and updating List of Tables...")
doc = Document('codenest_comprehensive.docx')

# Find the List of Tables section
paragraphs = doc.paragraphs
tables_index = None
for i, para in enumerate(paragraphs):
    if 'LIST OF TABLES' in para.text and para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        tables_index = i
        print(f"Found List of Tables at paragraph {i}")
        break

if not tables_index:
    print("List of Tables not found!")
    exit(1)

# Remove old table entries (next 20 paragraphs after the title)
print("Removing old table entries...")
for _ in range(20):
    if tables_index + 1 < len(paragraphs):
        p = paragraphs[tables_index + 1]
        p._element.getparent().remove(p._element)

# Now add updated table entries
print("Adding updated table entries...")

# Better, more relevant tables for CodeNest
tables = [
    ('Table 2.1', 'Comparison of Competitive Programming Platforms', '19'),
    ('Table 2.2', 'AI-Based vs Traditional Learning Systems', '21'),
    ('Table 2.3', 'Code Execution Security Approaches', '25'),
    ('Table 3.1', 'Existing System Limitations vs Proposed Solutions', '35'),
    ('Table 3.2', 'Feature Comparison: CodeNest vs Other Platforms', '38'),
    ('Table 4.1', 'Functional Requirements by Module', '47'),
    ('Table 4.2', 'Non-Functional Requirements Specifications', '50'),
    ('Table 4.3', 'Hardware Requirements (Development vs Production)', '52'),
    ('Table 4.4', 'Software Requirements and Versions', '54'),
    ('Table 4.5', 'Technology Stack Components', '56'),
    ('Table 5.1', 'Database Tables and Their Purposes', '63'),
    ('Table 5.2', 'Database Relationships and Cardinality', '64'),
    ('Table 5.3', 'Module Responsibilities and Dependencies', '68'),
    ('Table 5.4', 'API Endpoints by Category', '70'),
    ('Table 6.1', 'Programming Language Support Details', '85'),
    ('Table 6.2', 'Docker Container Resource Limits', '87'),
    ('Table 6.3', 'Security Measures Implementation', '91'),
    ('Table 7.1', 'Performance Test Results Summary', '107'),
    ('Table 7.2', 'User Testing Feedback Statistics', '111'),
    ('Table 7.3', 'Feature Comparison with Existing Platforms', '115'),
    ('Table 7.4', 'System Performance Metrics', '108'),
    ('Table 8.1', 'Challenges Encountered and Solutions', '123'),
    ('Table 8.2', 'Implementation Issues and Resolutions', '126'),
]

# Insert new table entries after the List of Tables heading
insert_position = tables_index + 2  # After heading and blank paragraph

for table_num, table_title, page_num in tables:
    p = doc.add_paragraph()
    p.add_run(f'{table_num}: {table_title}')
    p.add_run('\t' * 3)
    p.add_run(page_num)
    
    # Move this paragraph to the correct position
    p._element.getparent().remove(p._element)
    doc._element.body.insert(insert_position, p._element)
    insert_position += 1

doc.save('codenest_comprehensive.docx')

print("\n" + "="*70)
print("LIST OF TABLES UPDATED!")
print("="*70)
print(f"✓ Total tables: {len(tables)}")
print("\nUpdated tables are more relevant to CodeNest project:")
print("  - Platform comparisons")
print("  - System requirements")
print("  - Database design")
print("  - Performance metrics")
print("  - Testing results")
print("  - Challenges and solutions")
print("="*70)
